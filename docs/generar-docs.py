#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera los tres documentos de Funny en HTML, PDF y DOCX.

    python docs/generar-docs.py [--solo html|pdf|docx]

La fuente unica son los ficheros Markdown de `docs/`, que si van en commits. Lo
que sale de aqui NO se commitea: va a `docs/out/`, que esta en .gitignore, y se
publica como assets de la GitHub Release. Asi el repositorio sigue siendo ligero
de clonar y cada documento corresponde a una version concreta.

Por que no Pandoc, que es lo que pedia la plantilla: no esta instalado en el
entorno de desarrollo, y no se instala una herramienta a la ligera solo para
esto. Este script hace lo mismo con lo que si hay:

    Markdown -> HTML   markdown (Python), con la hoja de estilo de aqui abajo
    HTML     -> PDF    Microsoft Edge en modo headless (--print-to-pdf)
    Markdown -> DOCX   python-docx

Si falta alguna de las tres herramientas, el script lo dice y sigue con las que
puede, en lugar de fallar entero. Al final informa de que ha salido y que no.

Directiva propia: una linea

    <!-- incluir: OTRO.md -->

se sustituye por el contenido de ese fichero, quitandole su titulo de nivel 1 y
bajando un nivel el resto de sus encabezados. Es lo que permite que el manual
tecnico contenga el manual de usuario sin duplicar una sola linea.
"""
import argparse
import glob
import io
import os
import re
import shutil
import subprocess
import sys
import tempfile

AQUI = os.path.dirname(os.path.abspath(__file__))
SALIDA = os.path.join(AQUI, "out")

# Los ADR se leen en el repositorio y desde su indice, no tienen URL propia en
# el sitio, asi que no necesitan cabecera de Jekyll. El indice si.
SIN_CABECERA = {
    "adr/ADR-0001-stack.md",
    "adr/ADR-0002-sin-backend.md",
    "adr/ADR-0003-salon-nearby.md",
    "adr/ADR-0004-donacion-sin-facturacion.md",
    "adr/ADR-0005-contenido-sin-derechos.md",
    "adr/ADR-0006-licencia.md",
}

DOCUMENTOS = [
    ("MANUAL-USUARIO.md", "Funny — Manual de usuario"),
    ("MANUAL-TECNICO.md", "Funny — Manual tecnico"),
    ("GUIA-PUBLICACION.md", "Funny — Guia de publicacion"),
]

# --------------------------------------------------------------------------
# Hoja de estilo
#
# Un solo fichero HTML autocontenido, sin nada que descargar: los documentos
# tienen que poder leerse sin conexion y desde un ZIP. Los tamanos estan en
# puntos pensando en la impresion, que es de donde sale el PDF.
# --------------------------------------------------------------------------
ESTILO = """
:root {
  --tinta: #16161a;
  --tenue: #5a5a66;
  --linea: #d8d8e0;
  --fondo: #ffffff;
  --suave: #f6f6f9;
  --acento: #7a3cff;
}
* { box-sizing: border-box; }
body {
  margin: 0 auto; padding: 2.5rem 1.5rem 6rem; max-width: 46rem;
  font: 16px/1.65 -apple-system, "Segoe UI", Roboto, "Helvetica Neue", sans-serif;
  color: var(--tinta); background: var(--fondo);
  -webkit-text-size-adjust: 100%;
}
h1, h2, h3, h4 { line-height: 1.25; margin: 2.2em 0 .6em; font-weight: 650; }
h1 { font-size: 2.1rem; margin-top: 0; letter-spacing: -.02em; }
h2 { font-size: 1.5rem; padding-bottom: .3em; border-bottom: 2px solid var(--linea); }
h3 { font-size: 1.15rem; }
h4 { font-size: 1rem; color: var(--tenue); }
p, ul, ol, table, blockquote, pre { margin: 0 0 1.1em; }
a { color: var(--acento); text-decoration: none; border-bottom: 1px solid transparent; }
a:hover { border-bottom-color: var(--acento); }
code {
  font: .88em/1.5 ui-monospace, "Cascadia Code", Consolas, monospace;
  background: var(--suave); padding: .12em .35em; border-radius: 4px;
}
pre {
  background: var(--suave); padding: 1em 1.1em; border-radius: 8px;
  overflow-x: auto; border: 1px solid var(--linea);
}
pre code { background: none; padding: 0; font-size: .84em; line-height: 1.5; }
table { width: 100%; border-collapse: collapse; font-size: .93em; }
th, td { padding: .5em .7em; text-align: left; border-bottom: 1px solid var(--linea);
         vertical-align: top; }
th { background: var(--suave); font-weight: 650; }
blockquote {
  margin-left: 0; padding: .8em 1.1em; background: var(--suave);
  border-left: 4px solid var(--acento); border-radius: 0 6px 6px 0;
}
blockquote p:last-child { margin-bottom: 0; }
hr { border: 0; border-top: 1px solid var(--linea); margin: 2.5em 0; }
li { margin-bottom: .3em; }
li > ul, li > ol { margin-top: .3em; margin-bottom: .3em; }
.pie {
  margin-top: 4rem; padding-top: 1rem; border-top: 1px solid var(--linea);
  font-size: .85rem; color: var(--tenue);
}
input[type=checkbox] { margin-right: .4em; }

/*
  Ojo con este bloque: `page-break-after: avoid` en los encabezados JUNTO CON
  `page-break-inside: avoid` en tablas y bloques de codigo cuelga a Chromium al
  paginar el manual tecnico, que pasa de las 1.500 lineas. No falla ninguna de
  las tres reglas por separado —quitando cualquiera de ellas el PDF sale en tres
  segundos— sino su combinacion, asi que parece una explosion cuadratica en la
  paginacion. Se bisecto regla a regla para averiguarlo.

  La version de aqui abajo usa las propiedades modernas, deja que las tablas se
  partan (son largas a proposito) y controla las lineas huerfanas y viudas en
  lugar de prohibir el salto tras un encabezado. Un encabezado suelto al final de
  una pagina es un defecto cosmetico; un PDF que no se genera, no.
*/
@media print {
  @page { margin: 18mm 15mm; }
  body { max-width: none; padding: 0; font-size: 10.5pt; }
  h1 { font-size: 22pt; }
  h2 { font-size: 15pt; }
  h3 { font-size: 12pt; }
  pre, blockquote { break-inside: avoid; }
  p, li { orphans: 2; widows: 2; }
  a { color: var(--tinta); border: 0; }
  /* El destino de un enlace se pierde al imprimir, asi que se escribe. */
  a[href^="http"]::after { content: " (" attr(href) ")"; font-size: .8em; color: #666; }
}
"""


def leer(nombre):
    with io.open(os.path.join(AQUI, nombre), encoding="utf-8") as f:
        return sin_front_matter(f.read())


def sin_front_matter(texto):
    """Quita la cabecera YAML de Jekyll, si la hay.

    Los documentos la llevan porque sin ella GitHub Pages no los procesa y sus
    URL dan 404. Aqui no sirve de nada: sin quitarla, el `---` sale como una
    linea horizontal y el `title:` como un parrafo suelto al principio del PDF.
    """
    if not texto.startswith("---\n"):
        return texto
    fin = texto.find("\n---\n", 3)
    if fin == -1:
        return texto
    return texto[fin + len("\n---\n"):].lstrip("\n")


def resolver_includes(texto, vistos=None):
    """Sustituye `<!-- incluir: X.md -->` por X.md, bajando sus encabezados."""
    vistos = vistos or set()

    def uno(m):
        nombre = m.group(1).strip()
        if nombre in vistos:
            return "*(inclusion circular de %s omitida)*" % nombre
        contenido = resolver_includes(leer(nombre), vistos | {nombre})
        lineas = []
        for linea in contenido.split("\n"):
            # Fuera el titulo de nivel 1: el documento que incluye ya tiene el suyo.
            if linea.startswith("# "):
                continue
            # El resto baja un nivel, para que el indice siga teniendo sentido.
            if re.match(r"^#{2,5} ", linea):
                linea = "#" + linea
            lineas.append(linea)
        return "\n".join(lineas)

    return re.sub(r"^<!--\s*incluir:\s*(.+?)\s*-->\s*$", uno, texto, flags=re.M)


def version_del_proyecto():
    """Saca la version del build.gradle.kts, para no escribirla en dos sitios."""
    ruta = os.path.join(AQUI, "..", "app", "build.gradle.kts")
    try:
        with io.open(ruta, encoding="utf-8") as f:
            m = re.search(r'val versionSemVer = "([^"]+)"', f.read())
        return m.group(1) if m else "?"
    except OSError:
        return "?"


# ------------------------------------------------------------------ HTML

def a_html(md, titulo, version):
    try:
        import markdown
    except ImportError:
        return None
    cuerpo = markdown.markdown(
        md,
        extensions=["tables", "fenced_code", "toc", "sane_lists", "attr_list"],
        output_format="html5",
    )
    return (
        "<!DOCTYPE html>\n"
        '<html lang="es">\n<head>\n<meta charset="utf-8">\n'
        '<meta name="viewport" content="width=device-width, initial-scale=1">\n'
        "<title>%s</title>\n<style>%s</style>\n</head>\n<body>\n%s\n"
        '<div class="pie">Funny %s &middot; Brais Galdo &middot; Ghato Studio '
        "&middot; documento generado desde <code>docs/</code></div>\n"
        "</body>\n</html>\n" % (titulo, ESTILO, cuerpo, version)
    )


# ------------------------------------------------------------------- PDF

def buscar_edge():
    for ruta in (
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        "/usr/bin/microsoft-edge",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
    ):
        if os.path.exists(ruta):
            return ruta
    for nombre in ("microsoft-edge", "msedge", "google-chrome", "chromium"):
        encontrado = shutil.which(nombre)
        if encontrado:
            return encontrado
    return None


def a_pdf(ruta_html, ruta_pdf, navegador, espera=None):
    # El limite se calcula por tamano del HTML y no es fijo: el manual tecnico
    # incluye a los otros dos y pasa de las 1.500 lineas, y con 180 s fijos se
    # quedaba a medias mientras los cortos sobraban de tiempo.
    if espera is None:
        kb = os.path.getsize(ruta_html) / 1024.0
        espera = int(max(120, min(900, 60 + kb * 6)))
    # Perfil temporal: sin el, Edge reutiliza el del usuario y puede negarse a
    # arrancar en headless si ya hay una ventana abierta.
    perfil = tempfile.mkdtemp(prefix="funny-docs-")
    try:
        orden = [
            navegador,
            "--headless=new",
            "--disable-gpu",
            "--no-first-run",
            "--no-pdf-header-footer",
            "--user-data-dir=" + perfil,
            "--print-to-pdf=" + ruta_pdf,
            "file:///" + ruta_html.replace("\\", "/"),
        ]
        resultado = subprocess.run(
            orden, capture_output=True, timeout=espera, text=True, errors="replace"
        )
        if os.path.exists(ruta_pdf) and os.path.getsize(ruta_pdf) > 1024:
            return True, ""
        return False, (resultado.stderr or resultado.stdout or "sin salida").strip()[:400]
    except subprocess.TimeoutExpired:
        return False, "el navegador no ha terminado en %d s" % espera
    except OSError as e:
        return False, str(e)
    finally:
        shutil.rmtree(perfil, ignore_errors=True)


# ------------------------------------------------------------------ DOCX

def a_docx(md, titulo, ruta, version):
    """
    Markdown -> DOCX, con lo que hace falta y no mas.

    No es un conversor completo de Markdown: cubre encabezados, parrafos,
    listas, tablas, bloques de codigo, citas y reglas, que es todo lo que usan
    estos tres documentos. El formato en linea (negrita, cursiva, codigo,
    enlaces) se aplica parrafo a parrafo.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        return False, "falta python-docx"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading(titulo, 0)
    p = doc.add_paragraph("Funny %s · Brais Galdo · Ghato Studio" % version)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x66)
    doc.add_page_break()

    EN_LINEA = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*|\[.+?\]\(.+?\))")

    def escribir_en_linea(parrafo, texto):
        for trozo in EN_LINEA.split(texto):
            if not trozo:
                continue
            if trozo.startswith("**") and trozo.endswith("**"):
                parrafo.add_run(trozo[2:-2]).bold = True
            elif trozo.startswith("`") and trozo.endswith("`"):
                run = parrafo.add_run(trozo[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            elif trozo.startswith("*") and trozo.endswith("*") and len(trozo) > 2:
                parrafo.add_run(trozo[1:-1]).italic = True
            else:
                enlace = re.fullmatch(r"\[(.+?)\]\((.+?)\)", trozo)
                if enlace:
                    # Sin hipervinculo real: python-docx no lo trae de serie y
                    # montarlo a mano en el XML no merece la pena aqui.
                    parrafo.add_run(enlace.group(1)).underline = True
                else:
                    parrafo.add_run(trozo)

    def fila_de_tabla(linea):
        celdas = [c.strip() for c in linea.strip().strip("|").split("|")]
        return celdas

    lineas = md.split("\n")
    i = 0
    while i < len(lineas):
        linea = lineas[i]

        # Bloque de codigo
        if linea.startswith("```"):
            i += 1
            codigo = []
            while i < len(lineas) and not lineas[i].startswith("```"):
                codigo.append(lineas[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(codigo))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            i += 1
            continue

        # Tabla
        if linea.startswith("|") and i + 1 < len(lineas) and re.match(r"^\|[\s:|-]+\|?$", lineas[i + 1]):
            cabecera = fila_de_tabla(linea)
            i += 2
            filas = []
            while i < len(lineas) and lineas[i].startswith("|"):
                filas.append(fila_de_tabla(lineas[i]))
                i += 1
            tabla = doc.add_table(rows=1, cols=len(cabecera))
            tabla.style = "Light Grid Accent 1"
            for celda, texto in zip(tabla.rows[0].cells, cabecera):
                celda.text = ""
                escribir_en_linea(celda.paragraphs[0], texto)
                for run in celda.paragraphs[0].runs:
                    run.bold = True
            for fila in filas:
                celdas = tabla.add_row().cells
                for celda, texto in zip(celdas, fila):
                    celda.text = ""
                    escribir_en_linea(celda.paragraphs[0], texto)
            doc.add_paragraph()
            continue

        # Encabezado
        m = re.match(r"^(#{1,6})\s+(.*)$", linea)
        if m:
            nivel = min(len(m.group(1)), 4)
            texto = re.sub(r"[`*]", "", m.group(2))
            texto = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", texto)
            doc.add_heading(texto, nivel)
            i += 1
            continue

        # Regla
        if re.match(r"^-{3,}$", linea.strip()):
            doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Cita
        if linea.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            escribir_en_linea(p, linea[2:])
            i += 1
            continue

        # Lista
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", linea)
        if m:
            sangria = len(m.group(1)) // 2
            ordenada = m.group(2)[0].isdigit()
            estilo = "List Number" if ordenada else "List Bullet"
            if sangria:
                estilo += " %d" % min(sangria + 1, 3)
            texto = m.group(3)
            # Las casillas de verificacion del checklist, en texto.
            texto = re.sub(r"^\[ \]\s*", "☐ ", texto)
            texto = re.sub(r"^\[x\]\s*", "☑ ", texto, flags=re.I)
            try:
                p = doc.add_paragraph(style=estilo)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            escribir_en_linea(p, texto)
            i += 1
            continue

        # Parrafo
        if linea.strip():
            p = doc.add_paragraph()
            escribir_en_linea(p, linea)
        i += 1

    try:
        doc.save(ruta)
        return True, ""
    except OSError as e:
        return False, str(e)


# ------------------------------------------------------------------ main

def revisar_front_matter():
    """Comprueba la cabecera de Jekyll de cada documento de `docs/`.

    Existe porque un error aqui **no se ve hasta que GitHub Pages se cae**, y
    cuando se cae no publica nada: se queda la version anterior o un 404, y la
    politica de privacidad es una de las paginas afectadas. Paso de verdad: dos
    descripciones llevaban dos puntos seguidos de espacio sin entrecomillar
    —«no recoge ningun dato: la app...»— y para YAML eso es el separador de una
    clave, no texto. La compilacion entera fallo por dos lineas.

    La regla es mas estricta que YAML a proposito: **todo valor entrecomillado**.
    Asi no hay que acordarse de que caracteres son especiales (`:`, `#`, `&`,
    `*`, `[`, `{`, `|`, `>`, `%`, `@`) ni distinguir donde molestan.
    """
    problemas = []
    rutas = sorted(glob.glob(os.path.join(AQUI, "*.md")))
    rutas += sorted(glob.glob(os.path.join(AQUI, "adr", "*.md")))
    for ruta in rutas:
        with io.open(ruta, encoding="utf-8") as f:
            texto = f.read()
        nombre = os.path.relpath(ruta, AQUI).replace(os.sep, "/")
        if not texto.startswith("---\n"):
            # Los ADR no se sirven de uno en uno; solo su indice necesita cabecera.
            if nombre in SIN_CABECERA:
                continue
            problemas.append("%s: sin cabecera, Pages no lo procesara y dara 404" % nombre)
            continue
        fin = texto.find("\n---\n", 3)
        if fin == -1:
            problemas.append("%s: la cabecera no se cierra con ---" % nombre)
            continue
        for numero, linea in enumerate(texto[4:fin].split("\n"), start=2):
            if not linea.strip():
                continue
            m = re.match(r"^([a-z_]+): (.*)$", linea)
            if not m:
                problemas.append("%s:%d: no es «clave: valor» -> %s" % (nombre, numero, linea))
                continue
            valor = m.group(2)
            if not (valor.startswith('"') and valor.endswith('"') and len(valor) >= 2):
                problemas.append(
                    "%s:%d: el valor de «%s» tiene que ir entre comillas dobles"
                    % (nombre, numero, m.group(1))
                )
    return problemas


def main():
    ap = argparse.ArgumentParser(description="Genera los documentos de Funny.")
    ap.add_argument("--solo", choices=("html", "pdf", "docx"), help="un solo formato")
    args = ap.parse_args()

    problemas = revisar_front_matter()
    if problemas:
        print("La cabecera de Jekyll esta mal y GitHub Pages no compilaria:")
        for linea in problemas:
            print("   %s" % linea)
        return 1

    formatos = {args.solo} if args.solo else {"html", "pdf", "docx"}
    os.makedirs(SALIDA, exist_ok=True)
    version = version_del_proyecto()

    navegador = buscar_edge() if "pdf" in formatos else None
    if "pdf" in formatos and not navegador:
        print("AVISO: no se encuentra Edge ni Chrome; no habra PDF.")

    hechos, fallidos = [], []

    for nombre, titulo in DOCUMENTOS:
        base = os.path.splitext(nombre)[0]
        print("\n== %s" % nombre)
        md = resolver_includes(leer(nombre))
        print("   %d lineas tras resolver inclusiones" % md.count("\n"))

        ruta_html = os.path.join(SALIDA, base + ".html")
        if "html" in formatos or "pdf" in formatos:
            html = a_html(md, titulo, version)
            if html is None:
                fallidos.append((base + ".html", "falta el modulo markdown"))
            else:
                with io.open(ruta_html, "w", encoding="utf-8", newline="\n") as f:
                    f.write(html)
                if "html" in formatos:
                    hechos.append((base + ".html", os.path.getsize(ruta_html)))
                    print("   HTML  ok")

        if "pdf" in formatos and navegador and os.path.exists(ruta_html):
            ruta_pdf = os.path.join(SALIDA, base + ".pdf")
            ok, error = a_pdf(ruta_html, ruta_pdf, navegador)
            if ok:
                hechos.append((base + ".pdf", os.path.getsize(ruta_pdf)))
                print("   PDF   ok")
            else:
                fallidos.append((base + ".pdf", error))
                print("   PDF   FALLA: %s" % error)

        if "docx" in formatos:
            ruta_docx = os.path.join(SALIDA, base + ".docx")
            ok, error = a_docx(md, titulo, ruta_docx, version)
            if ok:
                hechos.append((base + ".docx", os.path.getsize(ruta_docx)))
                print("   DOCX  ok")
            else:
                fallidos.append((base + ".docx", error))
                print("   DOCX  FALLA: %s" % error)

        # El HTML solo hacia falta como paso intermedio del PDF.
        if "html" not in formatos and os.path.exists(ruta_html):
            os.remove(ruta_html)

    print("\n" + "=" * 62)
    print("Generados %d ficheros en docs/out/ (Funny %s)" % (len(hechos), version))
    for nombre, tam in hechos:
        print("   %-34s %8.1f KB" % (nombre, tam / 1024.0))
    if fallidos:
        print("\nNO generados (%d):" % len(fallidos))
        for nombre, error in fallidos:
            print("   %-34s %s" % (nombre, error))
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
